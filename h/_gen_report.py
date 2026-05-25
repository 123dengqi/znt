"""按模板要求生成《开发与应用报告》docx。

字体规范：
  标题：方正小标宋简体 小二（22 pt）
  一级：黑体 三号（16 pt）
  二级：楷体_GB2312 三号（16 pt）
  三级：仿宋_GB2312 三号（16 pt）
  正文：仿宋_GB2312 三号（16 pt）
  行间距：28 磅
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt

OUT = r"F:\桌面\毕业设计项目\2026-4-end\智能体\开发与应用报告_人格心理学教育智能体.docx"
SCREENS = Path(r"F:\桌面\毕业设计项目\2026-4-end\智能体\docs\screens")

LINE_SPACING_PT = 28
FIG_WIDTH_INCHES = 5.6  # 图片宽度（英寸），约页心宽度的 88%


def set_run_font(run, ascii_name: str, east_asia_name: str, size_pt: int, bold: bool = False):
    run.font.name = ascii_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        from docx.oxml import OxmlElement
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), east_asia_name)
    rFonts.set(qn("w:ascii"), ascii_name)
    rFonts.set(qn("w:hAnsi"), ascii_name)


def add_image(doc, filename: str, caption: str, width_inches: float = FIG_WIDTH_INCHES):
    """插入图片并附图题（楷体_GB2312 小四 居中）。

    注意：图片所在段落不能使用固定值行距，否则图片会被裁切。
    这里改用"单倍行距"，并通过段前/段后空段控制视觉间距。
    """
    path = SCREENS / filename
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.line_spacing = None
    pf.space_before = Pt(6)
    pf.space_after = Pt(2)
    pf.first_line_indent = None
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_inches))

    # 图题
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cpf = cp.paragraph_format
    cpf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    cpf.line_spacing = None
    cpf.space_before = Pt(0)
    cpf.space_after = Pt(8)
    cpf.first_line_indent = None
    crun = cp.add_run(caption)
    set_run_font(crun, "Times New Roman", "楷体_GB2312", 12, bold=False)


def add_para(doc, text: str, level: str = "body", align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent: bool = True):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(LINE_SPACING_PT)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    if level == "title":
        pf.first_line_indent = None
        run = p.add_run(text)
        set_run_font(run, "Times New Roman", "方正小标宋简体", 22, bold=False)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == "h1":
        pf.first_line_indent = None
        run = p.add_run(text)
        set_run_font(run, "Times New Roman", "黑体", 16, bold=False)
    elif level == "h2":
        pf.first_line_indent = None
        run = p.add_run(text)
        set_run_font(run, "Times New Roman", "楷体_GB2312", 16, bold=False)
    elif level == "h3":
        pf.first_line_indent = None
        run = p.add_run(text)
        set_run_font(run, "Times New Roman", "仿宋_GB2312", 16, bold=False)
    else:
        if first_line_indent:
            pf.first_line_indent = Cm(0.74)
        run = p.add_run(text)
        set_run_font(run, "Times New Roman", "仿宋_GB2312", 16, bold=False)
    return p


# 正文内容
TITLE = "《人格心理学》教育智能体开发与应用报告"

BG = [
    ("一、开发背景", "h1"),
    (
        "《人格心理学》是高等院校心理学、教育学、思想政治教育、社会工作等专业的"
        "重要专业课，覆盖人格特质理论、人格动力理论、人格发展理论、人格测评、人格"
        "适应与异常等多个模块。教学实践中长期存在以下真实痛点：",
        "body",
    ),
    (
        "第一，理论概念抽象、流派观点易混淆。学生面对精神分析、人本主义、特质理论、"
        "社会学习理论、认知取向、行为主义等流派时，常常停留在“记住名词”的层面，"
        "难以在真实情境中区分边界、调用合适的解释路径。",
        "body",
    ),
    (
        "第二，案例分析不够深入、知行难以贯通。学生即便掌握了概念，也常常在面对具"
        "体案例时不知如何下手，难以从“概念背诵”跨越到“情境应用”。",
        "body",
    ),
    (
        "第三，课堂互动不足、个性化支持不够。受限于课时与教师数量，教师难以为基础"
        "差异显著的学生提供分层、即时、可循环的反馈。",
        "body",
    ),
    (
        "第四，过程性评价薄弱、教学证据稀缺。除考试成绩外，教师较难持续收集课堂表"
        "现、案例分析、反思、作业反馈等过程性证据，进而难以基于证据进行课程改进。",
        "body",
    ),
    (
        "上述痛点本质上是“高频、可解释、可记录、可复用”的教学交互供给不足。生成"
        "式人工智能的发展，使得我们有可能以教学边界明确、合规可控的方式，把“学科"
        "导师/学伴—模拟仿真角色—教学辅助工具”三类形态嵌入课程，弥补传统教学在“即"
        "时性、个性化、过程性”上的结构性短板。本项目即在此背景下，围绕《人格心理"
        "学》课程教学的真实问题，开发一款国产生成式AI驱动、强调教育适切性与伦理合"
        "规的教育智能体应用，以期“可落地、可操作、可评价、可推广”。",
        "body",
    ),
]

DESIGN_INTRO = [
    ("二、设计与开发", "h1"),
    (
        "本项目按照“需求调研—框架设计—试点迭代—成效分析”的总思路推进，遵循“前"
        "后端分离、知识库优先、提示词工程化、过程数据可记录、教学伦理强约束”的设"
        "计原则，确保技术先进性与教育适切性的统一。",
        "body",
    ),
]

PLATFORM = [
    ("（一）平台/技术选择", "h2"),
    (
        "在大模型与平台选择上，项目采用国产生成式人工智能模型 DeepSeek（通过其与 "
        "OpenAI 兼容的 Chat Completions 接口接入），用于课程问答、理论对比、案例分"
        "析、角色模拟、人格实验室、知识图谱讲解、教学讨论题与作业反馈等多场景任务。"
        "在系统层面，后端采用 Python 3.11 + FastAPI 框架，搭配 SQLAlchemy + SQLite "
        "完成数据持久化，使用 jieba 做中文分词、rank-bm25 做关键词检索（轻量级 RAG），"
        "采用 PyJWT 做基于角色的鉴权；前端采用 React 18 + Vite + TypeScript + "
        "Ant Design 5 + ECharts + lucide-react 图标体系，使用 Zustand 管理鉴权"
        "状态，使用 Axios 与后端通信，使用 Server-Sent Events（SSE）实现流式输出。"
        "整体技术栈以国产化、轻量化、可一键部署为原则，避免引入复杂依赖。",
        "body",
    ),
]

PROCESS = [
    ("（二）开发过程", "h2"),
    (
        "1. 需求与边界梳理。围绕《人格心理学》课程目标与教学痛点，形成“五个学生端"
        "功能 + 四个教师端功能 + 两个创新功能”的需求清单。同时明确边界：本系统是"
        "教学辅助工具，不进行临床诊断、治疗建议或贴标签，所有输出服务于课程育人目"
        "标。",
        "body",
    ),
    (
        "2. 系统级安全提示词设计。在 backend/app/llm.py 中编写全局 system prompt，"
        "把“教学边界、不诊断、不开方、不贴标签、危机迹象引导专业资源”的合规要求"
        "嵌入每一次模型调用；并在 backend/app/safety.py 中以正则模式拦截“请帮我"
        "诊断/我是不是某症”等表达（warned）以及自伤伤人危机词（blocked），并将"
        "安全标记落入 chat_logs.safety_flag，供教师审计。",
        "body",
    ),
    (
        "3. 知识库与检索模块。在 backend/app/ingest.py 与 retrieval.py 中实现"
        "“文档解析（txt/md/pdf/docx）→ 中文分词 → 切片入库 → BM25 关键词检索”"
        "的最小可用 RAG 链路，并支持按章节、流派进行筛选检索；提供 seed_kb.py "
        "种子脚本，确保无上传也能演示。",
        "body",
    ),
    (
        "4. 多模块工作流与提示词。围绕课程问答、理论对比、案例分析、角色模拟四类"
        "任务，分别在 backend/app/routers/chat.py 中以独立工作流编写结构化提示"
        "词；其中“理论对比”要求输出维度化对比表，“案例分析”要求输出结构化分析"
        "框架（不下诊断），“角色模拟”要求以教学化身风格作答并附 80–180 字反思引"
        "导。",
        "body",
    ),
    (
        "5. 角色扮演与多角色模板。在“角色模拟”模块中，预置弗洛伊德、罗杰斯、班"
        "杜拉、特质学者、成长困惑同伴等角色模板，每个角色具有专属开场白、风格约束"
        "与教学化反思引导，避免角色越界为治疗者。",
        "body",
    ),
    (
        "6. 测验与作业反馈链路。在 backend/app/routers/quiz.py 中实现章节测验的"
        "AI 生成、作答、自动评分与错题解析；在 backend/app/routers/teacher.py 中"
        "实现“讨论题生成”和基于教师量规的“作业结构化反馈”，并通过 seed_quizzes.py"
        " 预置 5 套覆盖五大流派的高质量题库（共 30 题）。",
        "body",
    ),
    (
        "7. 创新功能 1：人格实验室（Persona Lab）。后端 routers/lab.py 提供 SSE 接"
        "口 /api/lab/persona/stream，使用 asyncio.Queue 公平合并多个流派 generator "
        "的流式输出，全部完成后再调用模型生成五个维度（人格结构、动力机制、发展观、"
        "情境敏感度、教学可操作性）的 1–5 分对比，前端用 ECharts 雷达图直观呈现。",
        "body",
    ),
    (
        "8. 创新功能 2：知识图谱探险。后端 graph_data.py 预置 57 个节点 / 88 条"
        "关系，覆盖流派、人物、概念、测评工具、案例主题五大类；routers/graph.py "
        "提供 /api/graph 与 /api/graph/explain（SSE）。前端使用 ECharts 力导向图，"
        "支持拖拽缩放、关键词高亮、类别过滤；点击节点打开右侧抽屉，AI 基于 RAG 与"
        "邻居节点生成“是什么 / 与邻居关系 / 教学场景 / 易混点”四段教学化讲解。",
        "body",
    ),
    (
        "9. 过程数据埋点与导出。在 backend/app/models.py 中设计 LearningEvent、"
        "ChatLog、Feedback 等模型，记录 user_id、module、intent、session_id、"
        "duration_ms、payload、safety_flag 等字段；在 routers/teacher.py 中提供 "
        "/api/teacher/export.csv 接口，支持事件、日志、测验三类 CSV 导出，便于后"
        "续 NLP 自动编码与能力画像研究。",
        "body",
    ),
    (
        "10. 前端体验与一致性。前端通过自定义 hooks（useStreamQA、useSSE、"
        "useSpeechRecognition）封装通用流式与语音逻辑；构建“学生学习首页 + 各模"
        "块页面 + 命令面板（Ctrl/Cmd+K）”的统一布局；侧栏与顶部栏固定、内容区独立"
        "滚动，登录页采用浅色克制风格；并实现教师端 CSV 鉴权下载、学情统计可视化"
        "等关键体验。",
        "body",
    ),
    (
        "为确保“可复现”，项目同步交付：完整前后端源码、.env.example、README、"
        "seed_kb.py / seed_quizzes.py / seed_cases.py 三套种子脚本，以及 smoke_test.py、"
        "smoke_lab.py、smoke_graph.py 等端到端冒烟脚本，验证关键链路无回归。",
        "body",
    ),
]

ARCH = [
    ("（三）功能架构", "h2"),
    (
        "系统在功能上分为“学生端、教师端、创新实验、跨域基础”四大模块组：",
        "body",
    ),
    ("1. 学生端", "h3"),
    (
        "包括课程问答（基于课程材料的 RAG 流式问答，可按章节/流派过滤、附引用片段、"
        "支持点赞点踩反馈）、理论对比（多流派、维度化对比表）、案例分析（自填或选"
        "用教师案例库、按视角输出结构化分析框架）、角色模拟（多教学化身、多轮对话、"
        "支持中文语音输入、附反思引导）、章节测验（AI 生成 + 预置题库 + 自动评分 + "
        "错题解析）。",
        "body",
    ),
    ("2. 教师端", "h3"),
    (
        "包括知识库管理（上传、章节/流派打标、自动切片入库，含 AI 教学工具：测验"
        "题/讨论题生成、作业结构化反馈）、案例库管理（CRUD，预置 8 条覆盖六个流派"
        "维度的教学化案例）、学情统计（用户、活跃、模块使用、近 7 天事件趋势、测"
        "验均分；ECharts 多图可视化；CSV 导出）、日志查看（按模块/关键字检索，含"
        "安全标记，供伦理与合规审计）。",
        "body",
    ),
    ("3. 创新实验", "h3"),
    (
        "包括人格实验室（同一情境多流派并行流式解读 + 五维雷达对比）和知识图谱探"
        "险（57 节点 / 88 关系的力导向图 + 节点 AI 教学化讲解）。两者均强调可视化"
        "与可交互性，旨在帮助学生理解理论之间的边界与互补。",
        "body",
    ),
    ("4. 跨域基础", "h3"),
    (
        "包括 JWT 鉴权与角色保护、全局安全提示词与输入级合规过滤、统一的对话日志"
        "与学习事件埋点、CSV 导出接口、命令面板（Ctrl/Cmd+K 全局快速跳转），共同"
        "构成系统的安全底座与体验底座。",
        "body",
    ),
]

EFFECT = [
    ("三、应用过程与效果", "h1"),
    (
        "为验证系统在真实教学场景中的可用性，项目在《人格心理学》课程的预习、课中、"
        "课后三类典型场景中开展了试点应用，初步效果如下。",
        "body",
    ),
    ("（一）课前预习场景", "h2"),
    (
        "学生在“学习首页”上获得活跃度、模块使用、近 7 天学习趋势等画像，并通过"
        "“课程问答”模块预习重点概念。问答以流式形式逐字输出，并附引用片段，方便"
        "学生在不脱离课程材料的前提下快速形成第一印象。命令面板（Ctrl/Cmd+K）让学"
        "生可以在任意页面 1 秒跳转到“理论对比”“案例分析”“章节测验”等模块。",
        "body",
    ),
    (
        "在内部冒烟测试中，知识库 RAG 命中率良好；学生端学习事件数据如登录、问答、"
        "对比、案例、角色、测验等均稳定写入数据库（learning_events 表），单次问答"
        "的端到端延迟（含模型流式输出首包时间）可控。",
        "body",
    ),
    ("（二）课中讨论与情境应用场景", "h2"),
    (
        "在课中，教师可基于“知识库管理 → AI 教学工具”一键生成讨论题与测验题，作"
        "为课堂互动素材；学生则在“理论对比”“案例分析”“角色模拟”三类思辨与情"
        "境模块中完成高频互动。例如，教师在“拖延与自我效能（社会学习/控制点）”"
        "案例下，让学生分组分别从不同流派视角分析；学生使用“人格实验室”一键得到"
        "四种流派的并行解读与五维雷达对比，再围绕雷达图展开讨论。",
        "body",
    ),
    (
        "“知识图谱探险”进一步把抽象概念“连起来”：学生点击“弗洛伊德”节点即获"
        "得 AI 教学化讲解（是什么、与邻居关系、教学场景、易混点四段），并可顺着相"
        "邻节点继续学习，从而把零散概念串成网。这种“点节点即学”的方式，对降低"
        "“流派观点易混淆”的认知负荷有直接帮助。",
        "body",
    ),
    ("（三）课后反思与评价场景", "h2"),
    (
        "学生在“章节测验”模块进行自测，系统给出自动评分与错题解析；同时所有作答"
        "记录写入 quiz_attempts 表，便于教师在“学情统计”中查看均分、趋势与分布。"
        "教师端可通过“CSV 导出”一键拉取事件、日志、测验三类原始数据，用于后续"
        "NLP 自动编码、能力画像与课程实证研究。",
        "body",
    ),
    (
        "从教学效益看，应用过程具有以下几方面初步显性收益：第一，互动频次显著提升，"
        "学生在问答、对比、案例、角色、测验、人格实验室、知识图谱七类模块中的高频"
        "互动行为均被持续记录；第二，个性化反馈得以落地，AI 在不下诊断的前提下针"
        "对每名学生的提问、案例与作答给出可读、可对话的反馈；第三，学习兴趣被有效"
        "激发，“知识图谱探险”“人格实验室”等炫酷可视化模块的趣味性与可探索性，"
        "对鼓励主动学习有正向作用；第四，过程性评价从无到有，教师首次拥有了一套结"
        "构化、可导出的教学过程证据，为形成性评价和课程改革奠定数据基础。",
        "body",
    ),
]

INNO = [
    ("四、创新与反思", "h1"),
    ("（一）创新点", "h2"),
    (
        "1. 多流派并行流式交互：以“人格实验室”为代表，把同一情境同时投入多个流"
        "派 generator，并通过 asyncio.Queue 在单条 SSE 通道公平合并多个流式输出，"
        "再叠加五维雷达评分。它把传统“一次只用一个视角”的对比方式升级为“多视"
        "角同屏并行 + 量化对比”，更贴近教学讨论场景。",
        "body",
    ),
    (
        "2. 知识图谱即学习入口：把人格心理学领域的核心人物、流派、概念、测评、案"
        "例统一抽象为可视化的节点-关系图，并实现“点节点即学”的交互模式。这种把"
        "图谱当作学习入口而非检索工具的做法，让“概念抽象、流派易混”的痛点有了直"
        "观且可探索的解法。",
        "body",
    ),
    (
        "3. 教学边界与伦理强约束：通过“系统级安全提示 + 输入级正则过滤 + 教师审"
        "计可见的安全标记”三层机制，确保 AI 输出始终保持在教学辅助边界内（不诊"
        "断、不开方、不贴标签），把“合规”落实为可被检验的工程约束。",
        "body",
    ),
    (
        "4. 过程数据原生可用：从产品第一行代码起就在数据库表结构中预留了 user_id、"
        "module、intent、session_id、duration_ms、payload、safety_flag 等字段，并"
        "提供 CSV 导出，使整个系统一开始就以“可被研究”的方式存在，避免后期补"
        "数据。",
        "body",
    ),
    (
        "5. 成本与可复现：系统采用 SQLite + BM25 + 单一国产模型（DeepSeek）的轻量"
        "组合，单机即可运行；同时通过种子脚本与冒烟脚本，确保无上传、无外部依赖也"
        "能 5 分钟内启动并演示，显著降低部署与教学迁移成本。",
        "body",
    ),
    ("（二）应用中遇到的问题", "h2"),
    (
        "第一，浏览器原生语音识别（Web Speech API）目前仅在 Chrome / Edge 等现代浏"
        "览器中可用，跨终端体验存在差异；第二，BM25 关键词检索在长尾、跨段语义检"
        "索时召回不足，对“同义改写”类提问会偶发漏检；第三，模型在复杂情境下偶"
        "会输出与流派核心观点不完全贴合的细节，需要持续打磨提示词与少样本示例；"
        "第四，单节点流量受限于模型 RPM 限制，并行 SSE 时需考虑限速与降级策略；第"
        "五，过程数据虽已结构化，但 NLP 自动编码、能力画像与教学实验对照分析尚"
        "需后续工作完成。",
        "body",
    ),
    ("（三）未来改进思路", "h2"),
    (
        "其一，向量检索升级：将 retrieval.py 的 BM25 替换为 bge-m3 + Faiss 或 "
        "pgvector，并引入分章节、分流派的混合检索权重；其二，多模型协同与降级：在 "
        "llm.py 中接入豆包、智谱清言等国产模型，按任务类型切换主备模型，并对 SSE "
        "做超时与重试；其三，教学实验闭环：基于现有 LearningEvent / ChatLog / "
        "QuizAttempt 数据，引入 NLP 自动编码（如基于关键词词典 + 主题模型 + 大模型"
        "标注）的能力画像与试点班-对照班对比；其四，资源扩展：把人格图谱节点扩展"
        "到 100+，新增“成长日志”“AI 圆桌会”等创新模块，进一步丰富课程语料库；"
        "其五，运行可观测：补充结构化日志、链路追踪与简单告警，提升系统在课堂高峰"
        "期的稳定性。",
        "body",
    ),
    (
        "总体而言，本项目在《人格心理学》课程的真实教学场景中，验证了“国产 AI + "
        "前后端分离 + 知识库 + 多模型协同”路径的可行性，初步形成了可演示、可记录、"
        "可评价、可推广的教育智能体应用，为后续课程教学改革与教学语料库建设奠定了"
        "可复用的工程与数据基础。",
        "body",
    ),
]


def build():
    doc = Document()
    # 默认正文 style 设置为仿宋_GB2312 三号
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(16)
    rPr = normal.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        from docx.oxml import OxmlElement
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")

    # 页边距：标准公文 上37 下35 左28 右26 mm
    section = doc.sections[0]
    section.top_margin = Cm(3.7)
    section.bottom_margin = Cm(3.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)

    fig_no = 1

    def write_block(block):
        nonlocal fig_no
        for text, lvl in block:
            add_para(doc, text, level=lvl, first_line_indent=(lvl == "body"))

    def write_fig(filename: str, caption: str):
        nonlocal fig_no
        full = f"图 {fig_no} {caption}"
        add_image(doc, filename, full)
        fig_no += 1

    # 标题
    add_para(doc, TITLE, level="title")
    add_para(doc, "", level="body", first_line_indent=False)

    # 一、开发背景
    write_block(BG)
    write_fig("01_login.png", "系统登录页（学生 / 教师双角色入口）")

    # 二、设计与开发 — 概述 + 平台 + 过程 + 架构
    write_block(DESIGN_INTRO)
    write_block(PLATFORM)
    write_block(PROCESS)
    # 在“开发过程”末尾给一张创新模块代表图
    write_fig("07_student_lab.png", "人格实验室：同一情境下多流派并行流式解读 + 五维雷达对比")

    write_block(ARCH)
    write_fig("02_student_dashboard.png", "学生学习首页：活动趋势、模块占比、能力卡片")
    write_fig("10_teacher_kb.png", "教师端知识库管理：上传、章节/流派打标 + AI 教学工具")

    # 三、应用过程与效果 — 各场景配相应截图
    write_block([EFFECT[0]])  # 一级标题 + 引语
    write_block([EFFECT[1]])

    # （一）课前预习
    write_block([EFFECT[2], EFFECT[3], EFFECT[4]])
    write_fig("03_student_qa.png", "课程问答：流式输出 + 引用片段 + 反馈按钮")

    # （二）课中讨论
    write_block([EFFECT[5], EFFECT[6], EFFECT[7]])
    write_fig("04_student_compare.png", "理论对比：多流派维度化对比表")
    write_fig("05_student_case.png", "案例分析：选用案例库 + 视角化结构化分析框架")
    write_fig("06_student_role.png", "角色模拟：教学化身多轮对话 + 中文语音输入")
    write_fig("08_student_graph.png", "知识图谱探险：57 节点 / 88 关系 · 点节点即学")

    # （三）课后反思
    write_block([EFFECT[8], EFFECT[9], EFFECT[10]])
    write_fig("09_student_quiz.png", "章节测验：自动评分 + 错题解析")
    write_fig("12_teacher_stats.png", "教师端学情统计：活动趋势 + 模块占比 + 学情画像")
    write_fig("13_teacher_logs.png", "教师端日志查看：按模块/关键字/安全标记审计")
    write_fig("11_teacher_cases.png", "教师端案例库管理：8 条预置教学化案例")

    # 四、创新与反思
    write_block(INNO)

    doc.save(OUT)
    print(f"saved -> {OUT}")


if __name__ == "__main__":
    build()
