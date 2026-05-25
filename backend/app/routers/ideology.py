"""课程思政融入设计中心 · 全部接口。

设计闭环：思政元素 → 知识点映射 → AI 教学设计生成 → 任务发布 → 学生反思
→ AI 反馈 → 教师评语 → 效果分析 / CSV 导出
"""
from __future__ import annotations

import csv
import io
import json
import re
import time
from collections import Counter
from datetime import datetime, timedelta
from typing import Any

import jieba
from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import StreamingResponse
from sqlalchemy import func
from sqlalchemy.orm import Session

from ..db import SessionLocal, get_db
from ..llm import LLMError, build_messages, chat_complete, chat_stream
from ..models import (
    IdeologyAnalyticsEvent,
    IdeologyElement,
    IdeologyReflection,
    IdeologyTeachingTask,
    KnowledgeIdeologyMapping,
    LearningEvent,
    User,
)
from ..schemas import (
    IdeologyElementIn,
    IdeologyElementOut,
    LessonPlanExportIn,
    LessonPlanIn,
    MappingGenerateIn,
    MappingIn,
    MappingOut,
    ReflectionIn,
    ReflectionOut,
    TaskIn,
    TaskOut,
    TeacherCommentIn,
)
from ..security import get_current_user, require_role

router = APIRouter(prefix="/api/ideology", tags=["ideology"])


# ---------------- 共用：思政专属 system prompt ----------------

IDEOLOGY_SYSTEM = (
    "你是《人格心理学》课程思政教学助手。你的任务是帮助教师把人格心理学知识点与课程思政元素自然融合。\n"
    "硬性约束：\n"
    "1) 输出内容必须体现专业知识、学科精神与价值导向的统一，避免空泛说教、避免口号化表达、避免脱离课程内容；\n"
    "2) 不得进行任何心理诊断、不得给学生贴标签、不得评价学生思想立场；\n"
    "3) 只能提供教学设计、案例任务、反思问题与学习反馈建议；\n"
    "4) 语言要学术、自然、克制，与《人格心理学》专业概念紧密结合。"
)

# ---------------- 工具函数 ----------------

def _now_event(db: Session, user_id: int, task_id: int | None, chapter: str, event_type: str, payload: dict | None = None):
    db.add(IdeologyAnalyticsEvent(
        user_id=user_id,
        task_id=task_id,
        chapter=chapter or "",
        event_type=event_type,
        payload=payload or {},
    ))


def _element_to_out(e: IdeologyElement) -> dict:
    return IdeologyElementOut.model_validate(e).model_dump()


def _mapping_to_out(m: KnowledgeIdeologyMapping, element_name: str) -> dict:
    d = MappingOut.model_validate(m).model_dump()
    d["element_name"] = element_name
    return d


def _task_to_out(t: IdeologyTeachingTask, submission_count: int = 0, submitted: bool = False) -> dict:
    d = TaskOut.model_validate(t).model_dump()
    d["submission_count"] = submission_count
    d["submitted"] = submitted
    return d


# ---------------- 思政元素库 ----------------

@router.get("/elements")
def list_elements(_: User = Depends(get_current_user), db: Session = Depends(get_db)):
    rows = db.query(IdeologyElement).order_by(IdeologyElement.id.desc()).all()
    return [_element_to_out(r) for r in rows]


@router.post("/elements")
def create_element(body: IdeologyElementIn, user: User = Depends(require_role("teacher", "admin")), db: Session = Depends(get_db)):
    e = IdeologyElement(
        name=body.name.strip(),
        category=body.category.strip(),
        description=body.description.strip(),
        suitable_chapters=body.suitable_chapters.strip(),
        example_cases=body.example_cases.strip(),
        created_by=user.id,
    )
    db.add(e)
    db.commit()
    db.refresh(e)
    return _element_to_out(e)


@router.put("/elements/{eid}")
def update_element(eid: int, body: IdeologyElementIn, _: User = Depends(require_role("teacher", "admin")), db: Session = Depends(get_db)):
    e = db.get(IdeologyElement, eid)
    if not e:
        raise HTTPException(status_code=404, detail="思政元素不存在")
    e.name = body.name.strip()
    e.category = body.category.strip()
    e.description = body.description.strip()
    e.suitable_chapters = body.suitable_chapters.strip()
    e.example_cases = body.example_cases.strip()
    db.commit()
    db.refresh(e)
    return _element_to_out(e)


@router.delete("/elements/{eid}")
def delete_element(eid: int, _: User = Depends(require_role("teacher", "admin")), db: Session = Depends(get_db)):
    e = db.get(IdeologyElement, eid)
    if not e:
        raise HTTPException(status_code=404, detail="思政元素不存在")
    db.delete(e)
    db.commit()
    return {"ok": True}


# ---------------- 知识点-思政元素映射 ----------------

@router.get("/mappings")
def list_mappings(
    chapter: str | None = None,
    ideology_element_id: int | None = None,
    _: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    q = db.query(KnowledgeIdeologyMapping)
    if chapter:
        q = q.filter(KnowledgeIdeologyMapping.chapter == chapter)
    if ideology_element_id is not None:
        q = q.filter(KnowledgeIdeologyMapping.ideology_element_id == ideology_element_id)
    rows = q.order_by(KnowledgeIdeologyMapping.id.desc()).all()
    elem_map = {e.id: e.name for e in db.query(IdeologyElement).all()}
    return [_mapping_to_out(r, elem_map.get(r.ideology_element_id, "")) for r in rows]


@router.post("/mappings")
def create_mapping(body: MappingIn, user: User = Depends(require_role("teacher", "admin")), db: Session = Depends(get_db)):
    if not db.get(IdeologyElement, body.ideology_element_id):
        raise HTTPException(status_code=400, detail="思政元素不存在")
    m = KnowledgeIdeologyMapping(
        chapter=body.chapter.strip(),
        knowledge_point=body.knowledge_point.strip(),
        ideology_element_id=body.ideology_element_id,
        integration_method=body.integration_method.strip(),
        teaching_scenario=body.teaching_scenario.strip(),
        prompt_template=body.prompt_template.strip(),
        evaluation_indicator=body.evaluation_indicator.strip(),
        created_by=user.id,
    )
    db.add(m)
    db.commit()
    db.refresh(m)
    elem = db.get(IdeologyElement, m.ideology_element_id)
    return _mapping_to_out(m, elem.name if elem else "")


@router.put("/mappings/{mid}")
def update_mapping(mid: int, body: MappingIn, _: User = Depends(require_role("teacher", "admin")), db: Session = Depends(get_db)):
    m = db.get(KnowledgeIdeologyMapping, mid)
    if not m:
        raise HTTPException(status_code=404, detail="映射不存在")
    m.chapter = body.chapter.strip()
    m.knowledge_point = body.knowledge_point.strip()
    m.ideology_element_id = body.ideology_element_id
    m.integration_method = body.integration_method.strip()
    m.teaching_scenario = body.teaching_scenario.strip()
    m.prompt_template = body.prompt_template.strip()
    m.evaluation_indicator = body.evaluation_indicator.strip()
    db.commit()
    db.refresh(m)
    elem = db.get(IdeologyElement, m.ideology_element_id)
    return _mapping_to_out(m, elem.name if elem else "")


@router.delete("/mappings/{mid}")
def delete_mapping(mid: int, _: User = Depends(require_role("teacher", "admin")), db: Session = Depends(get_db)):
    m = db.get(KnowledgeIdeologyMapping, mid)
    if not m:
        raise HTTPException(status_code=404, detail="映射不存在")
    db.delete(m)
    db.commit()
    return {"ok": True}


@router.post("/mappings/generate")
async def generate_mapping_suggestion(
    body: MappingGenerateIn,
    user: User = Depends(require_role("teacher", "admin")),
    db: Session = Depends(get_db),
):
    """AI 推荐：返回若干思政元素 + 融入方式建议（JSON）。"""
    elements = db.query(IdeologyElement).all()
    elem_list = "\n".join(f"- {e.name}（{e.category or '通用'}）：{e.description}" for e in elements)
    sys = (
        "请基于《人格心理学》课程，针对给定知识点，从下方思政元素库中挑选最贴切的 3 个，"
        "并给出'融入方式'建议（每个 1 句话，30-60 字，紧密结合人格心理学专业概念）。\n\n"
        f"【思政元素库】\n{elem_list or '（库为空，可自由命名 3 个常见元素）'}"
    )
    usr = (
        f"章节：{body.chapter or '（未填）'}\n"
        f"知识点：{body.knowledge_point}\n\n"
        "请严格输出 JSON，不要任何附加文字：\n"
        '{"suggestions":[{"element_name":"...","integration_method":"...","reason":"..."},...]}'
    )
    try:
        raw = await chat_complete(build_messages(sys, usr, extra_system=[IDEOLOGY_SYSTEM]), temperature=0.4, max_tokens=600)
    except LLMError as e:
        raise HTTPException(status_code=502, detail=str(e))

    text = raw.strip()
    if text.startswith("```"):
        text = text.strip("`")
        if text.lower().startswith("json"):
            text = text[4:]
    s, e2 = text.find("{"), text.rfind("}")
    suggestions: list[dict] = []
    if s != -1 and e2 != -1:
        try:
            data = json.loads(text[s : e2 + 1])
            suggestions = data.get("suggestions", []) or []
        except Exception:
            pass

    # 把 element_name 与现有库匹配，给 id（若有）
    name_to_id = {e.name: e.id for e in elements}
    for s_ in suggestions:
        s_["ideology_element_id"] = name_to_id.get(s_.get("element_name", ""), None)

    _now_event(db, user.id, None, body.chapter, "mapping_suggest", {"knowledge_point": body.knowledge_point, "n": len(suggestions)})
    db.commit()
    return {"suggestions": suggestions}


# ---------------- AI 教学设计生成 ----------------

def _lesson_plan_prompt(body: LessonPlanIn) -> tuple[str, str]:
    sys = (
        "请围绕《人格心理学》课程中的【{kp}】知识点，结合【{ie}】思政元素，"
        "生成一份{ts}的课程思政教学设计。\n"
        "要求严格输出 Markdown，并按以下章节展开：\n"
        "1) **专业知识目标**\n"
        "2) **思政育人目标**\n"
        "3) **融入切入点**\n"
        "4) **课前引导问题**（2-3 题，编号）\n"
        "5) **课中讨论任务**（2-3 题，编号）\n"
        "6) **课后反思问题**（2-3 题，编号）\n"
        "7) **案例材料**（1 个完整教学化案例，含人物、情境、关键冲突，120-200 字，避免真实身份）\n"
        "8) **评价指标**（4 个维度：专业关联度 / 表达完整性 / 伦理意识 / 反思深度，各给 1 句标准）\n"
        "9) **教师使用建议**（3 条要点）\n"
        "硬性约束：与人格心理学专业概念紧密结合，不下诊断、不贴标签、不评价学生思想立场、不口号化。"
    ).format(
        kp=body.knowledge_point,
        ie="、".join(body.ideology_elements) if body.ideology_elements else "（请自行选择 2 个最贴切的思政元素并在第一段说明）",
        ts=body.teaching_scenario,
    )
    usr = f"章节：{body.chapter}\n知识点：{body.knowledge_point}\n思政元素：{'、'.join(body.ideology_elements)}\n请开始撰写。"
    return sys, usr


@router.post("/lesson-plan/generate")
async def generate_lesson_plan(
    body: LessonPlanIn,
    user: User = Depends(require_role("teacher", "admin")),
):
    """SSE 流式生成教学设计。"""
    sys, usr = _lesson_plan_prompt(body)
    msgs = build_messages(sys, usr, extra_system=[IDEOLOGY_SYSTEM])

    async def gen():
        yield "data: " + json.dumps({"type": "meta", "knowledge_point": body.knowledge_point, "ideology_elements": body.ideology_elements}, ensure_ascii=False) + "\n\n"
        try:
            async for piece in chat_stream(msgs, temperature=0.5, max_tokens=1400):
                yield "data: " + json.dumps({"type": "delta", "text": piece}, ensure_ascii=False) + "\n\n"
        except LLMError as e:
            yield "data: " + json.dumps({"type": "error", "text": str(e)}, ensure_ascii=False) + "\n\n"
        yield "data: " + json.dumps({"type": "done"}, ensure_ascii=False) + "\n\n"

        db = SessionLocal()
        try:
            _now_event(db, user.id, None, body.chapter, "lesson_plan_generate",
                       {"knowledge_point": body.knowledge_point, "ideology_elements": body.ideology_elements})
            db.commit()
        finally:
            db.close()

    return StreamingResponse(gen(), media_type="text/event-stream", headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})


@router.post("/lesson-plan/export-docx")
def export_lesson_plan_docx(body: LessonPlanExportIn, _: User = Depends(require_role("teacher", "admin"))):
    """把 AI 生成的 Markdown 教学设计导出为 .docx（仿宋/宋体，标题加粗）。"""
    from io import BytesIO

    from docx import Document
    from docx.shared import Pt, Cm

    doc = Document()
    # 页边距
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(12)

    title = doc.add_heading(f"《人格心理学》课程思政教学设计 · {body.knowledge_point or '未命名'}", level=1)
    for run in title.runs:
        run.font.name = "黑体"

    meta = doc.add_paragraph()
    meta.add_run(f"章节：{body.chapter or '—'}    思政元素：{'、'.join(body.ideology_elements) if body.ideology_elements else '—'}")

    # 简单 Markdown 解析（标题/加粗/列表/正文）
    lines = (body.markdown or "").splitlines()
    for raw in lines:
        line = raw.rstrip()
        if not line.strip():
            doc.add_paragraph("")
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.lstrip().startswith(("- ", "* ", "• ")):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(line.lstrip()[2:])
        elif line.lstrip()[:2].rstrip(".").isdigit() and ". " in line:
            p = doc.add_paragraph(style="List Number")
            p.add_run(line.split(". ", 1)[1] if ". " in line else line)
        else:
            p = doc.add_paragraph()
            # 行内 **加粗**
            chunks = re.split(r"(\*\*.+?\*\*)", line)
            for c in chunks:
                if c.startswith("**") and c.endswith("**"):
                    r = p.add_run(c[2:-2])
                    r.bold = True
                else:
                    p.add_run(c)

    foot = doc.add_paragraph()
    foot.add_run("说明：本设计由 AI 协助生成，仅用于教学辅助，不进行临床诊断、不对学生贴标签。").italic = True

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    fname = f"思政教学设计-{body.knowledge_point or 'lesson'}.docx"
    # 简单 ASCII fallback 防止部分浏览器 Content-Disposition 解析异常
    from urllib.parse import quote
    return StreamingResponse(
        iter([buf.getvalue()]),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=lesson-plan.docx; filename*=UTF-8''{quote(fname)}"},
    )


# ---------------- 任务发布 / 列表 ----------------

@router.post("/tasks")
def create_task(body: TaskIn, user: User = Depends(require_role("teacher", "admin")), db: Session = Depends(get_db)):
    t = IdeologyTeachingTask(
        title=body.title.strip(),
        chapter=body.chapter.strip(),
        knowledge_point=body.knowledge_point.strip(),
        ideology_elements=body.ideology_elements.strip(),
        task_type=body.task_type.strip(),
        task_content=body.task_content,
        target_students=body.target_students.strip() or "all",
        status=body.status or "published",
        lesson_plan=body.lesson_plan or {},
        created_by=user.id,
    )
    db.add(t)
    db.commit()
    db.refresh(t)
    _now_event(db, user.id, t.id, t.chapter, "task_publish", {"title": t.title, "type": t.task_type})
    db.commit()
    return _task_to_out(t, 0, False)


@router.get("/tasks")
def list_tasks(
    chapter: str | None = None,
    status: str | None = "published",
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    q = db.query(IdeologyTeachingTask)
    if chapter:
        q = q.filter(IdeologyTeachingTask.chapter == chapter)
    if status:
        q = q.filter(IdeologyTeachingTask.status == status)
    rows = q.order_by(IdeologyTeachingTask.id.desc()).all()

    # 统计每个任务的提交数与当前学生是否已提交
    counts = dict(
        db.query(IdeologyReflection.task_id, func.count(IdeologyReflection.id))
        .group_by(IdeologyReflection.task_id)
        .all()
    )
    submitted_set: set[int] = set()
    if user.role == "student":
        submitted_set = set(
            tid for (tid,) in db.query(IdeologyReflection.task_id)
            .filter(IdeologyReflection.student_id == user.id)
            .distinct()
            .all()
        )

    out = []
    for t in rows:
        # 学生可见性过滤
        if user.role == "student":
            if t.status != "published":
                continue
            if t.target_students and t.target_students != "all":
                ids = {int(x) for x in t.target_students.split(",") if x.strip().isdigit()}
                if user.id not in ids:
                    continue
        out.append(_task_to_out(t, counts.get(t.id, 0), t.id in submitted_set))
    return out


@router.get("/tasks/{tid}")
def get_task(tid: int, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    t = db.get(IdeologyTeachingTask, tid)
    if not t:
        raise HTTPException(status_code=404, detail="任务不存在")
    sub_cnt = db.query(func.count(IdeologyReflection.id)).filter(IdeologyReflection.task_id == tid).scalar() or 0
    submitted = False
    if user.role == "student":
        submitted = db.query(IdeologyReflection.id).filter(
            IdeologyReflection.task_id == tid, IdeologyReflection.student_id == user.id
        ).first() is not None
        _now_event(db, user.id, tid, t.chapter, "task_view")
        db.commit()
    return _task_to_out(t, sub_cnt, submitted)


# ---------------- 学生反思提交（同步 AI 反馈） ----------------

REFLECTION_FEEDBACK_SYSTEM = (
    "你是《人格心理学》课程的教学反馈助手。学生针对一条思政教学任务提交了反思，"
    "请你只从以下四个维度给出教学性建议，每个维度 1-5 分（整数），并附 1 句改进建议：\n"
    "1) 专业关联度（与人格心理学专业概念结合的紧密程度）\n"
    "2) 表达完整性（结构是否清晰、要点是否齐全）\n"
    "3) 伦理意识（是否体现尊重个体、避免标签化、知情同意等专业伦理）\n"
    "4) 反思深度（是否有自我反观、举例迁移、可操作性）\n"
    "硬性约束：\n"
    "- 严禁评价学生'思想觉悟'、'立场对错'、'价值观水平'；\n"
    "- 严禁给学生贴标签、严禁做心理诊断或人格判断；\n"
    "- 只能给出教学性建议，语言克制、专业、温和；\n"
    "- 输出严格 JSON，结构：{\"dimensions\":{\"专业关联度\":x,\"表达完整性\":x,\"伦理意识\":x,\"反思深度\":x},"
    "\"comments\":{\"专业关联度\":\"...\",\"表达完整性\":\"...\",\"伦理意识\":\"...\",\"反思深度\":\"...\"},"
    "\"overall\":\"一段 80-140 字的整体性教学建议\"}"
)


async def _ai_feedback(task: IdeologyTeachingTask, text: str) -> tuple[str, dict]:
    """调用 LLM 生成反思反馈。返回 (overall_text, dimensions_dict)。"""
    usr = (
        f"【任务标题】{task.title}\n"
        f"【章节】{task.chapter}\n"
        f"【知识点】{task.knowledge_point}\n"
        f"【涉及思政元素】{task.ideology_elements}\n"
        f"【任务正文】{task.task_content[:600]}\n\n"
        f"【学生反思】{text[:1500]}\n\n"
        "请按系统约束严格输出 JSON。"
    )
    msgs = build_messages(REFLECTION_FEEDBACK_SYSTEM, usr, extra_system=[IDEOLOGY_SYSTEM])
    try:
        raw = await chat_complete(msgs, temperature=0.3, max_tokens=700)
    except LLMError:
        return ("（AI 反馈暂不可用，请稍后再试或联系教师）", {})

    t = raw.strip()
    if t.startswith("```"):
        t = t.strip("`")
        if t.lower().startswith("json"):
            t = t[4:]
    s, e = t.find("{"), t.rfind("}")
    if s == -1 or e == -1:
        return (raw[:600], {})
    try:
        data = json.loads(t[s : e + 1])
    except Exception:
        return (raw[:600], {})

    dims = data.get("dimensions", {}) or {}
    cmts = data.get("comments", {}) or {}
    overall = data.get("overall", "") or ""

    # 组合详细反馈文本
    lines = [overall.strip(), ""]
    lines.append("**分项建议**")
    for k in ["专业关联度", "表达完整性", "伦理意识", "反思深度"]:
        sc = dims.get(k, "-")
        ct = cmts.get(k, "")
        lines.append(f"- {k}（{sc}/5）：{ct}")
    feedback_text = "\n".join(lines).strip()
    return (feedback_text, {"scores": dims, "comments": cmts})


@router.post("/reflections")
async def submit_reflection(body: ReflectionIn, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    task = db.get(IdeologyTeachingTask, body.task_id)
    if not task:
        raise HTTPException(status_code=404, detail="任务不存在")
    if user.role != "student":
        raise HTTPException(status_code=403, detail="仅学生可提交反思")

    text = body.response_text.strip()
    if len(text) < 10:
        raise HTTPException(status_code=400, detail="反思内容过短，请至少撰写 10 字以上")

    t0 = time.time()
    ai_text, dims = await _ai_feedback(task, text)

    r = IdeologyReflection(
        task_id=task.id,
        student_id=user.id,
        response_text=text,
        ai_feedback=ai_text,
        ai_dimensions=dims,
        teacher_comment="",
        score=0.0,
    )
    db.add(r)
    db.flush()
    _now_event(db, user.id, task.id, task.chapter, "reflection_submit",
               {"length": len(text), "dimensions": dims.get("scores", {})})
    # 顺手写一条 LearningEvent 给学情统计使用
    db.add(LearningEvent(
        user_id=user.id,
        module="ideology",
        intent="reflection_submit",
        session_id="",
        duration_ms=int((time.time() - t0) * 1000),
        payload={"task_id": task.id, "chapter": task.chapter, "length": len(text)},
    ))
    db.commit()
    db.refresh(r)

    return ReflectionOut(
        id=r.id, task_id=r.task_id, student_id=r.student_id,
        response_text=r.response_text, ai_feedback=r.ai_feedback,
        ai_dimensions=r.ai_dimensions or {},
        teacher_comment=r.teacher_comment, score=r.score, submitted_at=r.submitted_at,
        task_title=task.title,
    ).model_dump()


@router.get("/reflections")
def list_reflections(
    task_id: int | None = None,
    student_id: int | None = None,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    q = db.query(IdeologyReflection)
    if user.role == "student":
        q = q.filter(IdeologyReflection.student_id == user.id)
    else:
        if student_id:
            q = q.filter(IdeologyReflection.student_id == student_id)
    if task_id:
        q = q.filter(IdeologyReflection.task_id == task_id)
    rows = q.order_by(IdeologyReflection.submitted_at.desc()).all()

    # 关联学生姓名 / 任务标题
    user_map = {u.id: (u.name or u.username) for u in db.query(User).all()}
    task_map = {t.id: t.title for t in db.query(IdeologyTeachingTask).all()}
    out = []
    for r in rows:
        d = ReflectionOut.model_validate(r).model_dump()
        d["student_name"] = user_map.get(r.student_id, "")
        d["task_title"] = task_map.get(r.task_id, "")
        out.append(d)
    return out


@router.put("/reflections/{rid}/comment")
def comment_reflection(rid: int, body: TeacherCommentIn, user: User = Depends(require_role("teacher", "admin")), db: Session = Depends(get_db)):
    r = db.get(IdeologyReflection, rid)
    if not r:
        raise HTTPException(status_code=404, detail="反思记录不存在")
    r.teacher_comment = body.teacher_comment.strip()
    r.score = max(0.0, min(100.0, float(body.score)))
    db.commit()
    db.refresh(r)
    _now_event(db, user.id, r.task_id, "", "teacher_comment", {"score": r.score})
    db.commit()
    return {"ok": True, "id": r.id, "score": r.score, "teacher_comment": r.teacher_comment}


# ---------------- 效果分析 ----------------

# 简易停用词（用于关键词词云）
_STOPWORDS = {
    "的", "了", "和", "是", "在", "我", "也", "都", "对", "与", "及", "或", "等", "就", "把",
    "让", "被", "更", "很", "会", "要", "并", "为", "从", "而", "这", "那", "之", "中", "上",
    "下", "他", "她", "它", "我们", "你们", "他们", "可以", "需要", "通过", "应该", "进行",
    "一个", "如何", "什么", "因为", "所以", "但是", "还有", "已经", "其实", "可能", "认为",
    "知识点", "思政", "学生", "教学", "课程", "心理学", "人格", "学习", "理论", "学科",
}


def _extract_keywords(texts: list[str], top_n: int = 30) -> list[dict]:
    if not texts:
        return []
    counter: Counter = Counter()
    for t in texts:
        for w in jieba.cut(t or ""):
            w = w.strip()
            if len(w) < 2 or w in _STOPWORDS:
                continue
            if not re.search(r"[\u4e00-\u9fffA-Za-z]", w):
                continue
            counter[w] += 1
    return [{"name": w, "value": c} for w, c in counter.most_common(top_n)]


@router.get("/analytics")
def analytics(_: User = Depends(require_role("teacher", "admin")), db: Session = Depends(get_db)):
    tasks = db.query(IdeologyTeachingTask).all()
    reflections = db.query(IdeologyReflection).all()
    students = db.query(User).filter(User.role == "student").all()
    student_count = len(students) or 1

    # 任务完成率（按 published 任务）
    pub_tasks = [t for t in tasks if t.status == "published"]
    task_completion = []
    for t in pub_tasks:
        if t.target_students and t.target_students != "all":
            ids = [int(x) for x in t.target_students.split(",") if x.strip().isdigit()]
            total = len(ids) or 1
        else:
            total = student_count
        submitted = sum(1 for r in reflections if r.task_id == t.id)
        rate = round(submitted / total * 100, 1) if total else 0
        task_completion.append({
            "task_id": t.id,
            "title": t.title,
            "chapter": t.chapter,
            "submitted": submitted,
            "total": total,
            "rate": rate,
        })

    # 章节完成率
    by_chapter: dict[str, dict] = {}
    for tc in task_completion:
        ch = tc["chapter"] or "未分章节"
        d = by_chapter.setdefault(ch, {"chapter": ch, "submitted": 0, "total": 0})
        d["submitted"] += tc["submitted"]
        d["total"] += tc["total"]
    chapter_completion = []
    for d in by_chapter.values():
        d["rate"] = round(d["submitted"] / d["total"] * 100, 1) if d["total"] else 0
        chapter_completion.append(d)
    chapter_completion.sort(key=lambda x: x["chapter"])

    # 思政元素关注度（任务 ideology_elements 字段 + 提交次数权重）
    elem_counter: Counter = Counter()
    for t in pub_tasks:
        names = [x.strip() for x in (t.ideology_elements or "").split(",") if x.strip()]
        # 该任务的提交数为该元素加权
        weight = sum(1 for r in reflections if r.task_id == t.id)
        for n in names:
            elem_counter[n] += 1 + weight
    element_distribution = [{"name": k, "value": v} for k, v in elem_counter.most_common(15)]

    # 提交趋势（近 14 天）
    today = datetime.utcnow().date()
    trend_days = [(today - timedelta(days=i)) for i in range(13, -1, -1)]
    by_day: Counter = Counter()
    for r in reflections:
        if r.submitted_at:
            by_day[r.submitted_at.date()] += 1
    submission_trend = [{"date": d.isoformat(), "value": by_day.get(d, 0)} for d in trend_days]

    # 高频关键词（学生反思）
    keywords = _extract_keywords([r.response_text for r in reflections], top_n=30)

    # 反思质量趋势（按提交时间序列，取 AI 平均分）
    quality_trend = []
    for r in sorted(reflections, key=lambda x: x.submitted_at):
        dims = (r.ai_dimensions or {}).get("scores", {}) if r.ai_dimensions else {}
        if not dims:
            continue
        try:
            avg = sum(float(dims.get(k, 0)) for k in ["专业关联度", "表达完整性", "伦理意识", "反思深度"]) / 4
        except Exception:
            continue
        quality_trend.append({"date": r.submitted_at.strftime("%m-%d %H:%M"), "value": round(avg, 2)})

    # 平均反思字数
    avg_len = round(sum(len(r.response_text) for r in reflections) / len(reflections), 1) if reflections else 0

    return {
        "summary": {
            "task_count": len(pub_tasks),
            "submission_count": len(reflections),
            "student_count": student_count,
            "avg_reflection_length": avg_len,
            "element_count": db.query(func.count(IdeologyElement.id)).scalar() or 0,
            "mapping_count": db.query(func.count(KnowledgeIdeologyMapping.id)).scalar() or 0,
        },
        "task_completion": task_completion,
        "chapter_completion": chapter_completion,
        "element_distribution": element_distribution,
        "submission_trend": submission_trend,
        "keywords": keywords,
        "quality_trend": quality_trend,
    }


# ---------------- 学生端 · 本章思政引导卡片 ----------------

@router.get("/student/guidance")
def student_guidance(chapter: str = Query(...), user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    """根据章节查询本章思政引导（首条映射 + 一条引导问题）。"""
    m = db.query(KnowledgeIdeologyMapping).filter(KnowledgeIdeologyMapping.chapter == chapter).order_by(KnowledgeIdeologyMapping.id.desc()).first()
    if not m:
        # 退化：找一条任务作为引导
        t = db.query(IdeologyTeachingTask).filter(IdeologyTeachingTask.chapter == chapter, IdeologyTeachingTask.status == "published").order_by(IdeologyTeachingTask.id.desc()).first()
        if not t:
            return {"available": False}
        return {
            "available": True,
            "chapter": chapter,
            "knowledge_point": t.knowledge_point or "",
            "ideology_elements": t.ideology_elements or "",
            "guidance_question": t.task_content[:120],
            "task_id": t.id,
        }
    elem = db.get(IdeologyElement, m.ideology_element_id)
    return {
        "available": True,
        "chapter": m.chapter,
        "knowledge_point": m.knowledge_point,
        "ideology_elements": elem.name if elem else "",
        "guidance_question": m.evaluation_indicator or m.integration_method or "请结合本章知识点谈谈你的理解。",
        "task_id": None,
    }


# ---------------- CSV 导出 ----------------

def _csv_response(headers: list[str], rows: list[list[Any]], filename: str) -> StreamingResponse:
    buf = io.StringIO()
    # BOM 让 Excel 直接识别 UTF-8
    buf.write("\ufeff")
    w = csv.writer(buf)
    w.writerow(headers)
    for r in rows:
        w.writerow(r)
    buf.seek(0)
    return StreamingResponse(iter([buf.getvalue()]), media_type="text/csv; charset=utf-8",
                             headers={"Content-Disposition": f"attachment; filename={filename}"})


@router.get("/export/csv")
def export_csv(
    kind: str = Query("reflections", regex="^(reflections|tasks|mappings|elements|events)$"),
    _: User = Depends(require_role("teacher", "admin")),
    db: Session = Depends(get_db),
):
    if kind == "reflections":
        rows = db.query(IdeologyReflection).order_by(IdeologyReflection.submitted_at.desc()).all()
        user_map = {u.id: (u.name or u.username) for u in db.query(User).all()}
        task_map = {t.id: t for t in db.query(IdeologyTeachingTask).all()}
        data = []
        for r in rows:
            t = task_map.get(r.task_id)
            dims = (r.ai_dimensions or {}).get("scores", {}) if r.ai_dimensions else {}
            data.append([
                r.id,
                r.task_id, (t.title if t else ""), (t.chapter if t else ""),
                r.student_id, user_map.get(r.student_id, ""),
                r.response_text.replace("\n", " "),
                r.ai_feedback.replace("\n", " "),
                dims.get("专业关联度", ""), dims.get("表达完整性", ""), dims.get("伦理意识", ""), dims.get("反思深度", ""),
                r.teacher_comment.replace("\n", " "),
                r.score,
                r.submitted_at.isoformat(),
            ])
        return _csv_response(
            ["id", "task_id", "task_title", "chapter", "student_id", "student_name", "response_text",
             "ai_feedback", "专业关联度", "表达完整性", "伦理意识", "反思深度", "teacher_comment", "score", "submitted_at"],
            data, "ideology_reflections.csv",
        )
    if kind == "tasks":
        rows = db.query(IdeologyTeachingTask).order_by(IdeologyTeachingTask.id.desc()).all()
        return _csv_response(
            ["id", "title", "chapter", "knowledge_point", "ideology_elements", "task_type", "status", "created_at"],
            [[t.id, t.title, t.chapter, t.knowledge_point, t.ideology_elements, t.task_type, t.status, t.created_at.isoformat()] for t in rows],
            "ideology_tasks.csv",
        )
    if kind == "mappings":
        rows = db.query(KnowledgeIdeologyMapping).order_by(KnowledgeIdeologyMapping.id.desc()).all()
        elem_map = {e.id: e.name for e in db.query(IdeologyElement).all()}
        return _csv_response(
            ["id", "chapter", "knowledge_point", "ideology_element", "integration_method", "teaching_scenario", "evaluation_indicator", "created_at"],
            [[m.id, m.chapter, m.knowledge_point, elem_map.get(m.ideology_element_id, ""), m.integration_method, m.teaching_scenario, m.evaluation_indicator, m.created_at.isoformat()] for m in rows],
            "ideology_mappings.csv",
        )
    if kind == "elements":
        rows = db.query(IdeologyElement).order_by(IdeologyElement.id.desc()).all()
        return _csv_response(
            ["id", "name", "category", "description", "suitable_chapters", "example_cases", "created_at"],
            [[e.id, e.name, e.category, e.description, e.suitable_chapters, e.example_cases, e.created_at.isoformat()] for e in rows],
            "ideology_elements.csv",
        )
    if kind == "events":
        rows = db.query(IdeologyAnalyticsEvent).order_by(IdeologyAnalyticsEvent.created_at.desc()).all()
        return _csv_response(
            ["id", "user_id", "task_id", "chapter", "event_type", "duration_ms", "payload", "created_at"],
            [[e.id, e.user_id, e.task_id, e.chapter, e.event_type, e.duration_ms, json.dumps(e.payload or {}, ensure_ascii=False), e.created_at.isoformat()] for e in rows],
            "ideology_events.csv",
        )
    raise HTTPException(status_code=400, detail="不支持的导出类型")
