from __future__ import annotations

import os
from typing import Tuple

from sqlalchemy.orm import Session

from .models import Document, Chunk
from .retrieval import chunk_text, tokenize


def _read_text_file(path: str) -> str:
    for enc in ("utf-8", "gbk", "utf-16"):
        try:
            with open(path, "r", encoding=enc) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    with open(path, "rb") as f:
        return f.read().decode("utf-8", errors="ignore")


def extract_text(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()
    if ext in (".txt", ".md", ".markdown"):
        return _read_text_file(path)
    if ext == ".pdf":
        try:
            from pypdf import PdfReader

            reader = PdfReader(path)
            return "\n".join((p.extract_text() or "") for p in reader.pages)
        except Exception as e:  # pragma: no cover
            return f"[PDF 解析失败: {e}]"
    if ext == ".docx":
        try:
            import docx  # python-docx

            d = docx.Document(path)
            return "\n".join(p.text for p in d.paragraphs)
        except Exception as e:  # pragma: no cover
            return f"[DOCX 解析失败: {e}]"
    return _read_text_file(path)


def ingest_document(
    db: Session,
    title: str,
    chapter: str,
    school: str,
    doc_type: str,
    file_path: str,
    user_id: int,
) -> Tuple[Document, int]:
    text = extract_text(file_path)
    doc = Document(
        title=title or os.path.basename(file_path),
        chapter=chapter or "",
        school=school or "",
        doc_type=doc_type or "material",
        filename=os.path.basename(file_path),
        uploaded_by=user_id,
    )
    db.add(doc)
    db.flush()
    chunks = chunk_text(text)
    n = 0
    for c in chunks:
        if not c.strip():
            continue
        toks = tokenize(c)
        db.add(
            Chunk(
                document_id=doc.id,
                chapter=doc.chapter,
                school=doc.school,
                text=c,
                tokens=" ".join(toks),
            )
        )
        n += 1
    db.commit()
    db.refresh(doc)
    return doc, n
